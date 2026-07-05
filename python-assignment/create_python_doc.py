from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Python Programming Assignment', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('GitHub Repository: https://github.com/urvish16/linux-assignment')
run.bold = True

doc.add_paragraph()

def add_code_block(doc, code):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)

def add_output_block(doc, output):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(output)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)

# Task 1
doc.add_heading('Task 1: Grade Checker', level=1)
doc.add_paragraph(
    'Takes a score as input and prints the corresponding grade using if-elif-else statements.'
)
doc.add_heading('Code:', level=2)
add_code_block(doc,
"""score = int(input("Enter your score: "))

if score >= 90:
    grade = "A"
elif score >= 80:
    grade = "B"
elif score >= 70:
    grade = "C"
elif score >= 60:
    grade = "D"
else:
    grade = "F"

print(f"Your grade is: {grade}")""")

doc.add_heading('Explanation:', level=2)
doc.add_paragraph(
    'The program uses if-elif-else to check the score range. '
    'It starts from the highest (90+) down to below 60, assigning grades A through F accordingly.'
)
doc.add_heading('Sample Output:', level=2)
add_output_block(doc,
"""Enter your score: 95  →  Your grade is: A
Enter your score: 75  →  Your grade is: C
Enter your score: 45  →  Your grade is: F""")

doc.add_paragraph()

# Task 2
doc.add_heading('Task 2: Student Grades (Dictionary)', level=1)
doc.add_paragraph(
    'A menu-driven program using a dictionary to store student names and grades. '
    'Supports adding, updating, and displaying student records.'
)
doc.add_heading('Code:', level=2)
add_code_block(doc,
"""students = {}

while True:
    print("\\n1. Add new student")
    print("2. Update existing student's grade")
    print("3. Print all student grades")
    print("4. Exit")

    choice = input("Enter your choice: ")

    if choice == "1":
        name = input("Enter student name: ")
        grade = input("Enter grade: ")
        students[name] = grade
        print(f"Student '{name}' added with grade '{grade}'.")

    elif choice == "2":
        name = input("Enter student name to update: ")
        if name in students:
            grade = input("Enter new grade: ")
            students[name] = grade
            print(f"Grade for '{name}' updated to '{grade}'.")
        else:
            print(f"Student '{name}' not found.")

    elif choice == "3":
        if students:
            print("\\nStudent Grades:")
            for name, grade in students.items():
                print(f"  {name}: {grade}")
        else:
            print("No students found.")

    elif choice == "4":
        print("Exiting...")
        break""")

doc.add_heading('Explanation:', level=2)
doc.add_paragraph(
    'A dictionary stores student name as key and grade as value. '
    'A while loop keeps the menu running until the user exits. '
    'Choice 1 adds a new entry, Choice 2 updates an existing one using in keyword check, '
    'Choice 3 loops through the dictionary using .items() to print all records.'
)
doc.add_heading('Sample Output:', level=2)
add_output_block(doc,
"""Student 'Alice' added with grade 'A'.
Student 'Bob' added with grade 'B'.
Grade for 'Alice' updated to 'A+'.

Student Grades:
  Alice: A+
  Bob: B""")

doc.add_paragraph()

# Task 3
doc.add_heading('Task 3: Write to a File', level=1)
doc.add_paragraph(
    'Creates a text file and writes content to it using Python\'s built-in file handling.'
)
doc.add_heading('Code:', level=2)
add_code_block(doc,
"""filename = "output.txt"

with open(filename, "w") as file:
    file.write("Hello, this is a Python file writing example.\\n")
    file.write("This is the second line.\\n")
    file.write("Python makes file handling easy!\\n")

print(f"Content written to '{filename}' successfully.")""")

doc.add_heading('Explanation:', level=2)
doc.add_paragraph(
    'open(filename, "w") opens the file in write mode, creating it if it does not exist. '
    'The with statement ensures the file is automatically closed after writing. '
    'file.write() writes each line of text to the file.'
)
doc.add_heading('Output:', level=2)
add_output_block(doc, "Content written to 'output.txt' successfully.")

doc.add_paragraph()

# Task 4
doc.add_heading('Task 4: Read from a File', level=1)
doc.add_paragraph(
    'Opens an existing text file in read mode and prints its contents to the screen.'
)
doc.add_heading('Code:', level=2)
add_code_block(doc,
"""filename = "output.txt"

with open(filename, "r") as file:
    content = file.read()

print("File contents:")
print(content)""")

doc.add_heading('Explanation:', level=2)
doc.add_paragraph(
    'open(filename, "r") opens the file in read mode. '
    'file.read() reads the entire file content as a string. '
    'The with statement ensures proper file closure after reading.'
)
doc.add_heading('Output:', level=2)
add_output_block(doc,
"""File contents:
Hello, this is a Python file writing example.
This is the second line.
Python makes file handling easy!""")

output_path = "/home/urvish-soni/Desktop/linux-assignment/python-assignment/Python_Assignment.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
